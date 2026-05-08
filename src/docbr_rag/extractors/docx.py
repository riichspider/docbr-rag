"""
Extrator de texto de documentos DOCX brasileiros.
Usa python-docx para parsear documentos Word.
"""

import re
from pathlib import Path
from typing import List, Tuple, Optional
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ..models import TipoDocumento


# Padrões específicos para DOCX brasileiros
_PADROES_TIPO_DOCX = {
    TipoDocumento.CONTRATO: [
        r"CONTRATO\s+(DE|SOCIAL|PARTICULAR|DE PRESTAÇÃO)",
        r"CONTRATANTE",
        r"CONTRATADO",
        r"CLÁUSULA\s+\w+",
        r"FORO",
        r"VIGÊNCIA",
    ],
    TipoDocumento.NFE: [
        r"NOTA\s+FISCAL\s+ELETR[ÔO]NICA",
        r"NF-?e",
        r"CHAVE\s+DE\s+ACESSO",
        r"DANFE",
        r"CNPJ",
    ],
    TipoDocumento.BOLETO: [
        r"BOLETO",
        r"BANCO\s+\w+\s+S\.?A\.?",
        r"C[ÓO]DIGO\s+DE\s+BARRAS",
        r"NOSSO\s+N[ÚU]MERO",
        r"VENCIMENTO",
    ],
    TipoDocumento.LAUDO: [
        r"LAUDO\s+(T[ÉE]CNICO|M[ÉE]DICO|PERICIAL)",
        r"PERITO",
        r"CONCLUS[ÃA]O",
        r"OBJETO\s+DA\s+PER[ÍI]CIA",
    ],
    TipoDocumento.CERTIDAO: [
        r"CERTID[ÃA]O",
        r"CERTIFIC[OA]",
        r"REGISTRO\s+CIVIL",
        r"CART[ÓO]RIO",
        r"MATR[ÍI]CULA",
    ],
    TipoDocumento.HOLERITE: [
        r"HOLERITE",
        r"CONTRACHEQUE",
        r"RECIBO\s+DE\s+PAGAMENTO",
        r"INSS",
        r"FGTS",
        r"SAL[ÁA]RIO\s+BASE",
    ],
}


def detectar_tipo_docx(texto: str) -> TipoDocumento:
    """
    Detecta o tipo de documento brasileiro com base no conteúdo DOCX.
    
    Args:
        texto: Texto extraído do documento DOCX
        
    Returns:
        TipoDocumento identificado ou DESCONHECIDO
    """
    texto_upper = texto.upper()
    pontuacao = {tipo: 0 for tipo in TipoDocumento}
    
    for tipo, padroes in _PADROES_TIPO_DOCX.items():
        for padrao in padroes:
            if re.search(padrao, texto_upper):
                pontuacao[tipo] += 1
    
    melhor_tipo = max(pontuacao, key=pontuacao.get)
    if pontuacao[melhor_tipo] == 0:
        return TipoDocumento.DESCONHECIDO
    
    return melhor_tipo


def extrair_texto_docx(caminho: str | Path) -> List[Tuple[int, str]]:
    """
    Extrai texto de um arquivo DOCX preservando estrutura.
    
    Args:
        caminho: Caminho para o arquivo DOCX
        
    Returns:
        Lista de tuplas (numero_secao, texto)
    """
    caminho = Path(caminho)
    if not caminho.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {caminho}")
    if not caminho.suffix.lower() == ".docx":
        raise ValueError(f"Arquivo deve ser DOCX: {caminho}")
    
    try:
        documento = docx.Document(caminho)
        secoes = []
        secao_atual = 1
        texto_secao = ""
        
        for paragrafo in documento.paragraphs:
            texto_paragrafo = paragrafo.text.strip()
            
            if not texto_paragrafo:
                continue
            
            # Detecta quebras de seção (títulos, cabeçalhos)
            if _eh_titulo_secao(texto_paragrafo):
                # Salva seção anterior se existir
                if texto_secao.strip():
                    secoes.append((secao_atual, texto_secao.strip()))
                    secao_atual += 1
                texto_secao = texto_paragrafo + "\n"
            else:
                texto_secao += texto_paragrafo + "\n"
        
        # Adiciona última seção
        if texto_secao.strip():
            secoes.append((secao_atual, texto_secao.strip()))
        
        # Se não houver seções claras, trata como documento único
        if not secoes:
            texto_completo = "\n".join([p.text for p in documento.paragraphs if p.text.strip()])
            secoes = [(1, texto_completo)]
        
        return [(i, _limpar_texto_docx(texto)) for i, texto in secoes]
        
    except Exception as e:
        raise RuntimeError(f"Não foi possível extrair texto do DOCX: {e}") from e


def extrair_metadados_docx(caminho: str | Path) -> dict:
    """
    Extrai metadados do documento DOCX.
    
    Args:
        caminho: Caminho para o arquivo DOCX
        
    Returns:
        Dicionário com metadados
    """
    caminho = Path(caminho)
    
    try:
        documento = docx.Document(caminho)
        metadados = {
            "arquivo": caminho.name,
            "tamanho": caminho.stat().st_size,
            "modificado": caminho.stat().st_mtime,
            "paragrafos": len(documento.paragraphs),
            "tabelas": len(documento.tables),
        }
        
        # Extrai propriedades do documento
        if documento.core_properties:
            props = documento.core_properties
            if props.author:
                metadados["autor"] = props.author
            if props.title:
                metadados["titulo"] = props.title
            if props.created:
                metadados["criado"] = props.created.isoformat()
            if props.modified:
                metadados["modificado_doc"] = props.modified.isoformat()
        
        # Detecta formatação especial
        metadados.update(_detectar_formatacao(documento))
        
        return metadados
        
    except Exception as e:
        return {"erro": f"Erro ao extrair metadados: {e}"}


def _eh_titulo_secao(texto: str) -> bool:
    """
    Detecta se um parágrafo é um título/seção.
    
    Args:
        texto: Texto do parágrafo
        
    Returns:
        True se for título/seção
    """
    # Padrões de títulos em documentos brasileiros
    padroes_titulo = [
        r"^CLÁUSULA\s+\w+",           # Cláusulas de contrato
        r"^ARTIGO\s+\d+",             # Artigos
        r"^§\s*\d+",                 # Parágrafos
        r"^\d+\.\d+\s+[A-Z]",       # Seções numeradas
        r"^[A-Z\s]{10,}$",          # Títulos em maiúsculas
        r"^CAP[ÍI]TULO\s+",          # Capítulos
        r"^SEÇÃO\s+",                # Seções explícitas
        r"^CONSIDERANDO",           # Considerandos
        r"^ONDE\s+",                # Onde (documentos legais)
        r"^PARTE\s+",               # Partes (contratos)
    ]
    
    texto_upper = texto.upper()
    return any(re.match(padrao, texto_upper) for padrao in padroes_titulo)


def _detectar_formatacao(documento) -> dict:
    """
    Detecta padrões de formatação no documento.
    
    Args:
        documento: Objeto Document do python-docx
        
    Returns:
        Dicionário com informações de formatação
    """
    formatacao = {
        "tem_negrito": False,
        "tem_italico": False,
        "tem_sublinhado": False,
        "alinhamentos": set(),
        "tem_tabelas": len(documento.tables) > 0,
        "tem_cabecalhos": False,
    }
    
    for paragrafo in documento.paragraphs:
        # Verifica alinhamento
        if paragrafo.alignment:
            if paragrafo.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                formatacao["alinhamentos"].add("centralizado")
            elif paragrafo.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                formatacao["alinhamentos"].add("direita")
            elif paragrafo.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                formatacao["alinhamentos"].add("justificado")
        
        # Verifica formatação de texto
        for run in paragrafo.runs:
            if run.bold:
                formatacao["tem_negrito"] = True
            if run.italic:
                formatacao["tem_italico"] = True
            if run.underline:
                formatacao["tem_sublinhado"] = True
    
    # Detecta cabeçalhos (parágrafos curtos e centralizados)
    for paragrafo in documento.paragraphs[:5]:  # Primeiros parágrafos
        if (len(paragrafo.text.strip()) < 100 and 
            paragrafo.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER):
            formatacao["tem_cabecalhos"] = True
            break
    
    # Converte set para lista
    formatacao["alinhamentos"] = list(formatacao["alinhamentos"])
    
    return formatacao


def _limpar_texto_docx(texto: str) -> str:
    """
    Limpa artefatos comuns em DOCX brasileiros.
    
    Args:
        texto: Texto bruto do DOCX
        
    Returns:
        Texto limpo
    """
    # Remove múltiplos espaços
    texto = re.sub(r" {2,}", " ", texto)
    
    # Remove múltiplas quebras de linha
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    
    # Remove tabs excessivos
    texto = re.sub(r"\t+", " ", texto)
    
    # Corrige aspas e caracteres especiais
    texto = texto.replace(""", '"')
    texto = texto.replace(""", '"')
    texto = texto.replace("'", "'")
    texto = texto.replace("'", "'")
    texto = texto.replace("…", "...")
    
    # Remove caracteres de controle
    texto = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", texto)
    
    return texto.strip()


def extrair_tabelas_docx(caminho: str | Path) -> List[List[List[str]]]:
    """
    Extrai tabelas de um documento DOCX.
    
    Args:
        caminho: Caminho para o arquivo DOCX
        
    Returns:
        Lista de tabelas (cada tabela é uma lista de linhas/células)
    """
    try:
        documento = docx.Document(caminho)
        tabelas_extraidas = []
        
        for tabela_idx, tabela in enumerate(documento.tables):
            tabela_dados = []
            
            for linha_idx, linha in enumerate(tabela.rows):
                linha_dados = []
                
                for celula in linha.cells:
                    texto_celula = celula.text.strip()
                    # Limpa formatação da célula
                    texto_celula = _limpar_texto_docx(texto_celula)
                    linha_dados.append(texto_celula)
                
                tabela_dados.append(linha_dados)
            
            tabelas_extraidas.append({
                "indice": tabela_idx,
                "linhas": len(tabela.rows),
                "colunas": len(tabela.columns) if tabela.rows else 0,
                "dados": tabela_dados
            })
        
        return tabelas_extraidas
        
    except Exception as e:
        raise RuntimeError(f"Erro ao extrair tabelas do DOCX: {e}") from e
